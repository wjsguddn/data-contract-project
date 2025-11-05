from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import logging
import json
import re

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
except ImportError:
    Document = None
    CT_Tbl = None
    CT_P = None
    logger.warning("python-docx가 설치되지 않았습니다. pip install python-docx")


"""
표준계약서 DOCX 파서
- 조(Article) / 항(Clause) / 호(SubClause) / 목(SubSubClause) / 별지(Exhibit) 구조 추출
- 표(Table) 구조화
- 장(Chapter)은 무시하고 조를 최상위 계층으로 처리
"""
class StdContractDocxParser:
    
    def __init__(self):
        if Document is None:
            raise ImportError("python-docx가 필요합니다: pip install python-docx")
    
    def parse(self, docx_path: Path, output_dir: Path) -> Dict[str, Path]:
        """
        DOCX 파싱 -> JSON 저장
        
        Args:
            docx_path: DOCX 파일 경로
            output_dir: 출력 디렉토리
            
        Return:
            저장된 파일 경로 딕셔너리
        """
        logger.info(f" DOCX 파싱 시작: {docx_path.name}")
        
        # 구조화 파싱
        structured_data = self.parse_contract_structure(docx_path)
        
        # 출력 디렉토리 생성
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # JSON 저장
        base_name = docx_path.stem
        output_file = output_dir / f"{base_name}_structured.json"
        
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(structured_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f" 구조화 파싱 완료: {output_file.name}")
        logger.info(f"   - 조(Article): {len(structured_data['articles'])}개")
        logger.info(f"   - 별지(Exhibit): {len(structured_data['exhibits'])}개")
        
        return {"structured": output_file}
    
    # ---------------------------
    # 표(Table) 및 문단 헬퍼 (docx 기반 경량 추출)
    # ---------------------------
    
    @staticmethod
    def get_first_meaningful_run(paragraph) -> Optional[Any]:
        """문단의 첫 의미있는 run 반환"""
        for r in paragraph.runs:
            if (r.text or "").strip():
                return r
        return paragraph.runs[0] if paragraph.runs else None
    
    @staticmethod
    def get_font_props(paragraph) -> Tuple[Optional[bool], Optional[float]]:
        """문단의 첫 의미있는 run 기준으로 bold, size_pt 추출"""
        run = StdContractDocxParser.get_first_meaningful_run(paragraph)
        if not run:
            return None, None
        size_pt = run.font.size.pt if run.font.size else None
        return run.bold, size_pt
    
    @staticmethod
    def paragraph_text(paragraph) -> str:
        """문단 텍스트 (앞 공백 보존)"""
        return paragraph.text or ""
    
    @staticmethod
    def extract_cell_text_docx(cell) -> str:
        """셀 텍스트 추출"""
        return (cell.text or "").strip()
    
    @staticmethod
    def is_cell_bold_docx(cell) -> bool:
        """셀이 bold인지 확인"""
        for para in cell.paragraphs:
            for run in para.runs:
                if (run.text or "").strip() and run.bold:
                    return True
        return False
    
    @staticmethod
    def is_all_bold_docx(cells: List[Any]) -> bool:
        """모든 셀이 bold인지 확인"""
        non_empty_cells = [c for c in cells if StdContractDocxParser.extract_cell_text_docx(c)]
        if not non_empty_cells:
            return False
        bold_count = sum(1 for c in non_empty_cells if StdContractDocxParser.is_cell_bold_docx(c))
        return bold_count == len(non_empty_cells)
    
    @staticmethod
    def detect_table_orientation_docx(table) -> str:
        """표 방향 감지 (row/column/none)"""
        rows = list(table.rows)
        if not rows or not rows[0].cells:
            return "none"
        first_row_all_bold = StdContractDocxParser.is_all_bold_docx(list(rows[0].cells))
        first_col_cells = [r.cells[0] for r in rows if r.cells]
        first_col_all_bold = StdContractDocxParser.is_all_bold_docx(first_col_cells)
        if first_row_all_bold and first_col_all_bold:
            return "row"
        if first_row_all_bold:
            return "row"
        if first_col_all_bold:
            return "column"
        return "none"
    
    @staticmethod
    def is_table_note_row_docx(row_cells: List[Any]) -> bool:
        """
        표의 행이 주석(note)인지 판단
        - 첫 번째 셀이 전체 열 병합 (grid_span >= 열 개수)
        - text alignment: RIGHT
        """
        if not row_cells:
            return False
        
        first_cell = row_cells[0]
        # 조건 1: 전체 병합 확인
        tcPr = first_cell._element.tcPr
        grid_span = None
        if tcPr is not None and tcPr.gridSpan is not None:
            grid_span = int(tcPr.gridSpan.val) if tcPr.gridSpan.val else 1
        if not grid_span or grid_span < len(row_cells):
            return False
        # 조건 2: 우측 정렬 확인
        if first_cell.paragraphs:
            align = first_cell.paragraphs[0].paragraph_format.alignment
            align_str = str(align) if align is not None else None
            if align_str and align_str.startswith("RIGHT"):
                return True
        
        return False
    
    @staticmethod
    def parse_table_simple_docx(table) -> Dict[str, Any]:
        """
        테이블 구조 변환 (병합 무시, 중복 허용)
        
        모든 셀을 독립적으로 처리
        마지막 행이 주석인 경우 notes 필드로 분리
        
        반환:
        {
            "type": "표",
            "orientation": "row" | "column",
            "headers": [...],
            "data": [{"열": "값"}, ...],
            "notes": "..." (선택적)
        }
        """
        rows = list(table.rows)
        if not rows:
            return {
                "type": "표",
                "orientation": "none",
                "headers": [],
                "data": []
            }
        
        # 표 방향 감지
        orientation = StdContractDocxParser.detect_table_orientation_docx(table)
        
        # 마지막 행이 주석인지 확인
        notes = None
        data_rows = rows[1:]  # 헤더 제외 (row 기반 처리 기본)
        
        if len(data_rows) > 0 and StdContractDocxParser.is_table_note_row_docx(list(data_rows[-1].cells)):
            # 마지막 행을 주석으로 분리
            notes = StdContractDocxParser.extract_cell_text_docx(data_rows[-1].cells[0])
            data_rows = data_rows[:-1]
        
        if orientation == "row":
            # 1행이 헤더 (병합 무시, 중복 제거)
            raw_headers = [StdContractDocxParser.extract_cell_text_docx(c) for c in rows[0].cells]
            
            # 중복된 헤더 고유 키(_n) 생성
            headers = []
            header_counts = {}
            for h in raw_headers:
                if h in header_counts:
                    header_counts[h] += 1
                    headers.append(f"{h}_{header_counts[h]}")
                else:
                    header_counts[h] = 0
                    headers.append(h)
            
            # 데이터 행 파싱
            data = []
            for row in data_rows:
                row_dict = {}
                for i, cell in enumerate(row.cells):
                    header = headers[i] if i < len(headers) else f"열{i+1}"
                    row_dict[header] = StdContractDocxParser.extract_cell_text_docx(cell)
                data.append(row_dict)
        
        elif orientation == "column":
            # 1열이 헤더
            headers = [StdContractDocxParser.extract_cell_text_docx(r.cells[0]) for r in rows if r.cells]
            data = []
            
            num_cols = len(rows[0].cells) if rows else 0
            for col_idx in range(1, num_cols):
                col_dict = {}
                for row_idx, r in enumerate(rows):
                    if col_idx < len(r.cells):
                        header = headers[row_idx] if row_idx < len(headers) else f"항목{row_idx+1}"
                        col_dict[header] = StdContractDocxParser.extract_cell_text_docx(r.cells[col_idx])
                if col_dict:
                    data.append(col_dict)
        
        else:
            # 헤더 없음
            headers = [f"열{i+1}" for i in range(len(rows[0].cells))]
            data = []
            for r in rows:
                row_dict = {}
                for i, cell in enumerate(r.cells):
                    header = headers[i] if i < len(headers) else f"열{i+1}"
                    row_dict[header] = StdContractDocxParser.extract_cell_text_docx(cell)
                data.append(row_dict)
        
        result = {
            "type": "표",
            "orientation": orientation,
            "headers": headers,
            "data": data
        }
        
        # notes가 있으면 추가
        if notes:
            result["notes"] = notes
        
        return result
    
    # ---------------------------
    # 패턴 매칭 함수 (docx Paragraph 기반)
    # ---------------------------
    
    @staticmethod
    def is_chapter(para) -> Optional[str]:
        """장(Chapter) 판별:
        - "제n장"으로 시작
        - bold=true
        - size_pt=12.0
        """
        text = (StdContractDocxParser.paragraph_text(para)).strip()
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        if bold is True and size_pt == 12.0:
            match = re.match(r'^제(\d)장', text)
            if match:
                return text
        return None
    
    @staticmethod
    def is_article(para) -> Optional[Tuple[int, str]]:
        """
        조(Article) 판별:
        - "제n조"으로 시작
        - bold=true
        - size_pt=11.0
        반환: (조 번호, 조 전체 텍스트) 또는 None
        """
        text = (StdContractDocxParser.paragraph_text(para)).strip()
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        if bold is True and size_pt == 11.0:
            # "제n조"으로 시작하는지 체크 (n은 1~2자리)
            match = re.match(r'^제(\d{1,2})조', text)
            if match:
                article_num = int(match.group(1))
                return article_num, text
        
        return None
    
    @staticmethod
    def is_clause(para) -> Optional[Tuple[int, str]]:
        """
        항(Clause) 판별:
        - "  ①"으로 시작
        - bold=null 또는 true
        - size_pt=11.0
        반환: (항 번호, 항 전체 텍스트) 또는 None
        """
        text = StdContractDocxParser.paragraph_text(para)  # 앞 공백 보존
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        # bold는 null 또는 true 모두 허용
        if size_pt == 11.0:
            # "  ①" 형태로 시작하는지 체크 (공백 2개 + 원 숫자)
            # 원문자: ①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮ ...
            match = re.match(r'^  ([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳])', text)
            if match:
                # 원 숫자를 실제 숫자로 변환
                circle_nums = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳'
                clause_num = circle_nums.index(match.group(1)) + 1
                return clause_num, text
        
        return None
    
    @staticmethod
    def is_subclause(para) -> Optional[Tuple[int, str]]:
        """
        호(SubClause) 판별:
        - "  n." (조 하위) 또는 "    n." (항 하위)으로 시작
        - bold=null
        - size_pt=11.0
        반환: (호 번호, 호 전체 텍스트) 또는 None
        """
        text = StdContractDocxParser.paragraph_text(para)
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        # bold=null
        if bold is None and size_pt == 11.0:
            # 문단 시작 "  n." 또는 "    n." 형태 체크 (공백 2개 또는 4개)
            match = re.match(r'^  {1,2}(\d{1,2})\.', text)
            if match:
                subclause_num = int(match.group(1))
                return subclause_num, text
        
        return None
    
    @staticmethod
    def is_subsubclause(para) -> Optional[Tuple[int, str]]:
        """
        목(SubSubClause) 판별:
        - "    가." (호 하위) 또는 "      가." (더 깊은 들여쓰기) 등으로 시작
        - bold=null
        - size_pt=11.0
        반환: (목 번호, 목 전체 텍스트) 또는 None
        
        한글 순서: 가나다라마바사아자차카타파하
        """
        text = StdContractDocxParser.paragraph_text(para)
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        # bold=null
        if bold is None and size_pt == 11.0:
            # "    가." 또는 "      가." 형태 체크 (공백 4개 이상 + 한글 1자 + 점)
            match = re.match(r'^    +([가-힣])\.', text)
            if match:
                # 한글을 숫자로 변환
                korean_order = '가나다라마바사아자차카타파하'
                korean_char = match.group(1)
                if korean_char in korean_order:
                    subsubclause_num = korean_order.index(korean_char) + 1
                    return subsubclause_num, text
        
        return None
    
    @staticmethod
    def is_exhibit_index(para) -> Optional[Tuple[int, str]]:
        """
        별지 색인(ExhibitIndex) 판별:
        - "[별지n]"으로 시작
        - bold=true
        - size_pt=11.0
        반환: (별지 번호, 별지 전체 텍스트) 또는 None
        """
        text = (StdContractDocxParser.paragraph_text(para)).strip()
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        if bold is True and size_pt == 11.0:
            # "[별지n]"으로 시작하는지 체크
            match = re.match(r'^\[별지(\d)\]', text)
            if match:
                exhibit_num = int(match.group(1))
                return exhibit_num, text
        
        return None
    
    @staticmethod
    def is_article_text(para) -> bool:
        """
        본문(ArticleText) 판별:
        - 다른 항목 패턴에 해당하지 않는 일반 텍스트
        - bold=null 또는 true
        - size_pt=11.0
        """
        text = (StdContractDocxParser.paragraph_text(para)).strip()
        if not text:
            return False
        
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        if size_pt == 11.0:
            # 다른 패턴에 해당하지 않는지 확인
            if (not StdContractDocxParser.is_article(para) and 
                not StdContractDocxParser.is_clause(para) and 
                not StdContractDocxParser.is_subclause(para) and 
                not StdContractDocxParser.is_exhibit_index(para)):
                return True
        
        return False
    
    @staticmethod
    def is_exhibit_content(para) -> bool:
        """
        별지 내용(Exhibit) 판별:
        - 별지범위 내 구역(계층범위 외)
        - size_pt<=11.0 or None
        """
        text = (StdContractDocxParser.paragraph_text(para)).strip()
        if not text:
            return False
        
        bold, size_pt = StdContractDocxParser.get_font_props(para)
        
        # size_pt가 10.0~11.0 범위이거나 None (스타일 상속)
        if size_pt is None or (size_pt <= 11.0):
            return True
        
        return False
    
    # ---------------------------
    # 계층 구조 빌더
    # ---------------------------
    
    @staticmethod
    def new_chapter(text: str) -> Dict[str, Any]:
        """장(Chapter) 노드 생성"""
        return {
            "type": "장",
            "text": text,
            "articles": []
        }
    
    @staticmethod
    def new_article(number: int, text: str) -> Dict[str, Any]:
        """조(Article) 노드 생성"""
        return {
            "type": "조",
            "number": number,
            "text": text,
            "content": []  # 조 본문, 항, table 등이 순서대로 들어감
        }
    
    @staticmethod
    def new_clause(number: int, text: str) -> Dict[str, Any]:
        """항(Clause) 노드 생성"""
        return {
            "type": "항",
            "number": number,
            "text": text,
            "content": []  # 항 본문, 호, table 등이 순서대로 들어감
        }
    
    @staticmethod
    def new_subclause(number: int, text: str) -> Dict[str, Any]:
        """호(SubClause) 노드 생성"""
        return {
            "type": "호",
            "number": number,
            "text": text,
            "content": []  # 목(SubSubClause)이 들어갈 수 있음
        }
    
    @staticmethod
    def new_subsubclause(number: int, text: str) -> Dict[str, Any]:
        """목(SubSubClause) 노드 생성"""
        return {
            "type": "목",
            "number": number,
            "text": text
        }
    
    @staticmethod
    def new_article_text(text: str) -> Dict[str, Any]:
        """본문(ArticleText) 노드 생성"""
        return {
            "type": "조 본문",
            "text": text
        }
    
    @staticmethod
    def new_clause_text(text: str) -> Dict[str, Any]:
        """항 본문(ClauseText) 노드 생성"""
        return {
            "type": "항 본문",
            "text": text
        }
    
    @staticmethod
    def new_exhibit(number: int, title: str) -> Dict[str, Any]:
        """별지(Exhibit) 노드 생성"""
        return {
            "type": "별지",
            "number": number,
            "title": title,
            "content": []
        }
    
    @staticmethod
    def new_exhibit_text(text: str) -> Dict[str, Any]:
        """별지 본문 노드 생성"""
        return {
            "type": "별지 본문",
            "text": text
        }
    
    # ---------------------------
    # 메인 파서 (docx 스트리밍)
    # ---------------------------
    
    def parse_contract_structure(self, docx_path: Path) -> Dict[str, Any]:
        """
        계약서 구조 파싱
        
        반환 구조:
        {
            "articles": [  # 조가 최상위 계층 (장은 무시)
                {
                    "type": "조",
                    "number": 1,
                    "text": "제1조...",
                    "content": [  # 조 본문, 항, 호, table 등을 등장 순서대로 추가
                        {"type": "조 본문", "text": "..."},
                        {
                            "type": "항",
                            "number": 1,
                            "text": "  ① ...",
                            "content": [  # 항 본문, 호, 목, table 등을 등장 순서대로
                                {"type": "항 본문", "text": "..."},
                                {
                                    "type": "호",
                                    "number": 1,
                                    "text": "    1. ...",
                                    "content": [  # 목이 있는 경우
                                        {"type": "목", "number": 1, "text": "    가. ..."}
                                    ]
                                }
                            ]
                        },
                        {
                            "type": "호",  # 조 바로 하위에도 호 가능
                            "number": 1,
                            "text": "    1. ...",
                            "content": []
                        }
                    ]
                }
            ],
            "exhibits": [
                {
                    "type": "별지",
                    "number": 1,
                    "title": "[별지1] ...",
                    "content": [...]
                }
            ]
        }
        """
        doc = Document(str(docx_path))
        result = {
            "articles": [],
            "exhibits": []
        }
        
        # 상태 변수
        current_article: Optional[Dict[str, Any]] = None
        current_clause: Optional[Dict[str, Any]] = None
        current_subclause: Optional[Dict[str, Any]] = None
        in_exhibit_section = False  # 별지 섹션에 진입했는지
        current_exhibit: Optional[Dict[str, Any]] = None
        
        for element in doc.element.body:
            # 표 처리
            if isinstance(element, CT_Tbl):
                # 해당 Table 객체 찾기
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                if table:
                    parsed_table = self.parse_table_simple_docx(table)
                    if in_exhibit_section and current_exhibit:
                        current_exhibit["content"].append(parsed_table)
                    elif current_clause:
                        current_clause["content"].append(parsed_table)
                    elif current_article:
                        current_article["content"].append(parsed_table)
                continue

            # 문단 처리
            if isinstance(element, CT_P):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para is None:
                    continue
            
            # 별지 색인 체크 (우선)
            exhibit_match = self.is_exhibit_index(para)
            if exhibit_match:
                in_exhibit_section = True
                exhibit_num, exhibit_title = exhibit_match
                current_exhibit = self.new_exhibit(exhibit_num, exhibit_title)
                result["exhibits"].append(current_exhibit)
                # 별지 섹션에서는 조/항/호 컨텍스트 초기화
                current_article = None
                current_clause = None
                current_subclause = None
                continue
            
            # 별지 섹션 내부인 경우
            if in_exhibit_section:
                if current_exhibit and self.is_exhibit_content(para):
                    # 별지 본문으로 변환
                    current_exhibit["content"].append(self.new_exhibit_text(self.paragraph_text(para)))
                continue
            
            # Chapter 체크 - 무시 (장은 건너뛰고 조를 최상위로 처리)
            chapter_text = self.is_chapter(para)
            if chapter_text:
                # 장은 무시하고 다음 요소로 진행
                continue
            
            # Article 체크
            article_match = self.is_article(para)
            if article_match:
                article_num, article_text = article_match
                current_article = self.new_article(article_num, article_text)
                
                # 조를 최상위 articles에 직접 추가
                result["articles"].append(current_article)
                
                current_clause = None
                current_subclause = None
                continue
            
            # Clause 체크
            clause_match = self.is_clause(para)
            if clause_match:
                if current_article:
                    clause_num, clause_text = clause_match
                    current_clause = self.new_clause(clause_num, clause_text)
                    # 조의 content에 항을 순서대로 추가
                    current_article["content"].append(current_clause)
                continue
            
            # SubClause 체크
            subclause_match = self.is_subclause(para)
            if subclause_match:
                subclause_num, subclause_text = subclause_match
                current_subclause = self.new_subclause(subclause_num, subclause_text)
                
                if current_clause:
                    # 항의 content에 호를 순서대로 추가
                    current_clause["content"].append(current_subclause)
                elif current_article:
                    # 조 바로 하위에도 호 추가 가능 (항 없이)
                    current_article["content"].append(current_subclause)
                continue
            
            # SubSubClause 체크 (목)
            subsubclause_match = self.is_subsubclause(para)
            if subsubclause_match:
                if current_subclause:
                    # 호의 content에 목을 순서대로 추가
                    subsubclause_num, subsubclause_text = subsubclause_match
                    subsubclause_node = self.new_subsubclause(subsubclause_num, subsubclause_text)
                    current_subclause["content"].append(subsubclause_node)
                continue
            
            # Article(Clause)Text 체크
            if self.is_article_text(para):
                if current_clause:
                    # clause_text
                    current_clause["content"].append(self.new_clause_text(self.paragraph_text(para)))
                elif current_article:
                    # article_text
                    current_article["content"].append(self.new_article_text(self.paragraph_text(para)))
                continue
        
        return result